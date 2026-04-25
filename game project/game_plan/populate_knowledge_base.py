import os
from langchain_community.vectorstores import Chroma

# 本地嵌入模型模拟类（避免模型下载）
class LocalEmbeddings:
    def __init__(self):
        # 模拟嵌入模型，返回固定长度的向量
        self.vector_length = 384
    
    def embed_documents(self, texts):
        # 为每个文本返回一个固定长度的向量
        return [[0.0 for _ in range(self.vector_length)] for _ in texts]
    
    def embed_query(self, text):
        # 为查询返回一个固定长度的向量
        return [0.0 for _ in range(self.vector_length)]

def load_documents_from_folder(folder_path):
    """从文件夹中加载所有支持的文档"""
    documents = []
    
    # 遍历文件夹
    for root, dirs, files in os.walk(folder_path):
        for file in files:
            # 跳过 Chroma 数据库文件
            if file.endswith('.bin') or file.endswith('.pkl') or file.endswith('.sqlite3'):
                continue
                
            file_path = os.path.join(root, file)
            file_ext = os.path.splitext(file)[1].lower()
            
            # 检查文件类型是否支持
            try:
                if file_ext == '.txt':
                    # 加载文本文件
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        content = f.read()
                    documents.append({'page_content': content, 'metadata': {'source': file_path}})
                    print(f"成功加载文件: {file_path}")
                    
                elif file_ext == '.pdf':
                    # 加载 PDF 文件
                    from pypdf import PdfReader
                    reader = PdfReader(file_path)
                    for page_num in range(len(reader.pages)):
                        page = reader.pages[page_num]
                        content = page.extract_text()
                        if content:
                            documents.append({'page_content': content, 'metadata': {'source': file_path, 'page': page_num + 1}})
                    print(f"成功加载文件: {file_path}")
                    
                elif file_ext == '.docx':
                    # 加载 Word 文件
                    from docx import Document
                    doc = Document(file_path)
                    content = '\n'.join([para.text for para in doc.paragraphs])
                    if content:
                        documents.append({'page_content': content, 'metadata': {'source': file_path}})
                    print(f"成功加载文件: {file_path}")
                    
                else:
                    print(f"跳过不支持的文件类型: {file_path}")
                    
            except Exception as e:
                print(f"加载文件失败 {file_path}: {str(e)}")
    
    return documents

def main():
    try:
        # 初始化本地嵌入模型模拟
        embeddings = LocalEmbeddings()
        
        # 文件夹路径（素材和知识库都在同一个目录）
        folder_path = "C:/Users/Lenovo/Desktop/游戏ai agent知识库"
        
        # 加载文档
        print(f"开始加载文件夹中的文档: {folder_path}")
        documents = load_documents_from_folder(folder_path)
        
        if not documents:
            print("没有加载到任何文档")
            return
        
        # 分割文档（简单分割，避免依赖 langchain-text-splitters）
        split_docs = []
        for doc in documents:
            # 简单按 1000 字符分割
            content = doc['page_content']
            chunks = [content[i:i+1000] for i in range(0, len(content), 800)]  # 800 是重叠部分
            for chunk in chunks:
                split_docs.append({'page_content': chunk, 'metadata': doc['metadata']})
        
        print(f"成功分割 {len(documents)} 个文档为 {len(split_docs)} 个块")
        
        # 转换为 LangChain 文档格式
        try:
            from langchain_core.documents import Document as LangChainDocument
            langchain_docs = [LangChainDocument(page_content=doc['page_content'], metadata=doc['metadata']) for doc in split_docs]
            print(f"成功转换为 LangChain 文档格式")
        except Exception as e:
            print(f"转换文档格式失败: {str(e)}")
            return
        
        # 加载或创建知识库
        try:
            vectorstore = Chroma(
                persist_directory=folder_path,
                embedding_function=embeddings
            )
            print(f"成功加载或创建知识库")
        except Exception as e:
            print(f"加载或创建知识库失败: {str(e)}")
            return
        
        # 添加文档
        try:
            vectorstore.add_documents(langchain_docs)
            print(f"成功添加 {len(langchain_docs)} 条文档到知识库")
            print(f"知识库中文档数量: {vectorstore._collection.count()}")
        except Exception as e:
            print(f"添加文档失败: {str(e)}")
            
    except Exception as e:
        print(f"运行失败: {str(e)}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main()